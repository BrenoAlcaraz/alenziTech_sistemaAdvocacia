from pathlib import Path


class ErroImportacaoDocumento(Exception):
    pass


def extrair_texto_documento(arquivo):
    """Extrai texto de um arquivo PDF ou DOCX e devolve string limpa."""
    nome = getattr(arquivo, "name", "") or ""
    extensao = Path(nome).suffix.lower()

    if extensao == ".pdf":
        return _extrair_pdf(arquivo)
    elif extensao == ".docx":
        return _extrair_docx(arquivo)
    else:
        raise ErroImportacaoDocumento("Envie um arquivo PDF ou DOCX.")


def _extrair_pdf(arquivo):
    try:
        from pypdf import PdfReader
        from pypdf.errors import PdfReadError
    except ImportError:
        raise ErroImportacaoDocumento("Biblioteca de leitura de PDF não disponível.")

    try:
        leitor = PdfReader(arquivo)
    except Exception:
        raise ErroImportacaoDocumento(
            "Não foi possível abrir o PDF. O arquivo pode estar corrompido ou protegido."
        )

    if leitor.is_encrypted:
        raise ErroImportacaoDocumento(
            "O PDF está protegido por senha e não pode ser importado."
        )

    textos = []
    for pagina in leitor.pages:
        try:
            texto = pagina.extract_text() or ""
        except Exception:
            texto = ""
        if texto.strip():
            textos.append(texto.strip())

    conteudo = "\n\n".join(textos).strip()

    if not conteudo:
        raise ErroImportacaoDocumento(
            "Não foi possível extrair texto deste PDF. "
            "Ele pode ser um documento escaneado."
        )

    return conteudo


def _extrair_docx(arquivo):
    try:
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
    except ImportError:
        raise ErroImportacaoDocumento("Biblioteca de leitura de DOCX não disponível.")

    try:
        documento = Document(arquivo)
    except Exception:
        raise ErroImportacaoDocumento(
            "Não foi possível abrir o arquivo DOCX. "
            "O arquivo pode estar corrompido ou em formato inválido."
        )

    linhas = [p.text for p in documento.paragraphs if p.text.strip()]

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                texto_celula = celula.text.strip()
                if texto_celula:
                    linhas.append(texto_celula)

    conteudo = "\n".join(linhas).strip()

    if not conteudo:
        raise ErroImportacaoDocumento(
            "Não foi possível encontrar texto reutilizável neste documento."
        )

    return conteudo
